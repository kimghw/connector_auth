"""Tests for file converters."""

import unittest
from pathlib import Path
import tempfile
import os

from ..converters import (
    PDFConverter,
    DOCXConverter,
    HWPConverter,
    ExcelConverter,
    OCRConverter
)


class TestPDFConverter(unittest.TestCase):
    """Test PDF converter."""

    def setUp(self):
        """Setup test fixtures."""
        self.converter = PDFConverter()
        self.test_dir = Path(__file__).parent / 'sample_files'

    def test_supports(self):
        """Test file support detection."""
        self.assertTrue(self.converter.supports('test.pdf'))
        self.assertTrue(self.converter.supports('test.PDF'))
        self.assertFalse(self.converter.supports('test.txt'))
        self.assertFalse(self.converter.supports('test.docx'))

    def test_convert_simple_pdf(self):
        """Test converting a simple PDF."""
        # Create a simple test PDF if pdfplumber is available
        try:
            import pdfplumber
            # Test with actual PDF file if available
            pdf_path = self.test_dir / 'test.pdf'
            if pdf_path.exists():
                text = self.converter.convert(str(pdf_path))
                self.assertIsInstance(text, str)
                self.assertTrue(len(text) > 0)
        except ImportError:
            self.skipTest("pdfplumber not installed")


class TestDOCXConverter(unittest.TestCase):
    """Test DOCX converter."""

    def setUp(self):
        """Setup test fixtures."""
        self.converter = DOCXConverter()

    def test_supports(self):
        """Test file support detection."""
        self.assertTrue(self.converter.supports('test.docx'))
        self.assertTrue(self.converter.supports('test.doc'))
        self.assertTrue(self.converter.supports('test.DOCX'))
        self.assertFalse(self.converter.supports('test.pdf'))
        self.assertFalse(self.converter.supports('test.txt'))

    def test_convert_simple_docx(self):
        """Test converting a simple DOCX."""
        try:
            from docx import Document

            # Create a temporary DOCX file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                doc = Document()
                doc.add_paragraph("Test paragraph 1")
                doc.add_paragraph("Test paragraph 2")
                doc.save(tmp.name)

                # Convert
                text = self.converter.convert(tmp.name)
                self.assertIn("Test paragraph 1", text)
                self.assertIn("Test paragraph 2", text)

                # Cleanup
                os.unlink(tmp.name)

        except ImportError:
            self.skipTest("python-docx not installed")


class TestExcelConverter(unittest.TestCase):
    """Test Excel converter."""

    def setUp(self):
        """Setup test fixtures."""
        self.converter = ExcelConverter()

    def test_supports(self):
        """Test file support detection."""
        self.assertTrue(self.converter.supports('test.xlsx'))
        self.assertTrue(self.converter.supports('test.xls'))
        self.assertTrue(self.converter.supports('test.xlsm'))
        self.assertFalse(self.converter.supports('test.pdf'))
        self.assertFalse(self.converter.supports('test.csv'))

    def test_convert_simple_excel(self):
        """Test converting a simple Excel file."""
        try:
            import openpyxl

            # Create a temporary Excel file
            with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                workbook = openpyxl.Workbook()
                sheet = workbook.active
                sheet.title = "Test Sheet"

                # Add some data
                sheet['A1'] = 'Name'
                sheet['B1'] = 'Value'
                sheet['A2'] = 'Test'
                sheet['B2'] = 123

                workbook.save(tmp.name)
                workbook.close()

                # Convert
                text = self.converter.convert(tmp.name)
                self.assertIn("Test Sheet", text)
                self.assertIn("Name", text)
                self.assertIn("Value", text)
                self.assertIn("Test", text)
                self.assertIn("123", text)

                # Cleanup
                os.unlink(tmp.name)

        except ImportError:
            self.skipTest("openpyxl not installed")


class TestOCRConverter(unittest.TestCase):
    """Test OCR converter."""

    def setUp(self):
        """Setup test fixtures."""
        self.converter = OCRConverter()

    def test_supports(self):
        """Test file support detection."""
        self.assertTrue(self.converter.supports('test.png'))
        self.assertTrue(self.converter.supports('test.jpg'))
        self.assertTrue(self.converter.supports('test.jpeg'))
        self.assertTrue(self.converter.supports('test.gif'))
        self.assertTrue(self.converter.supports('test.bmp'))
        self.assertFalse(self.converter.supports('test.pdf'))
        self.assertFalse(self.converter.supports('test.txt'))

    def test_convert_simple_image(self):
        """Test converting a simple image."""
        try:
            import pytesseract
            from PIL import Image, ImageDraw, ImageFont

            # Create a simple test image with text
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                # Create image with text
                img = Image.new('RGB', (200, 100), color='white')
                draw = ImageDraw.Draw(img)

                # Try to use a basic font
                try:
                    # This might fail on some systems
                    font = ImageFont.load_default()
                    draw.text((10, 10), "Test Text", fill='black', font=font)
                except:
                    # Fallback to no font specification
                    draw.text((10, 10), "Test Text", fill='black')

                img.save(tmp.name)

                # Convert
                text = self.converter.convert(tmp.name)
                self.assertIsInstance(text, str)
                # OCR might not perfectly recognize the text, so just check it's not empty
                self.assertTrue(len(text.strip()) > 0)

                # Cleanup
                os.unlink(tmp.name)

        except ImportError:
            self.skipTest("pytesseract or Pillow not installed")


if __name__ == '__main__':
    unittest.main()