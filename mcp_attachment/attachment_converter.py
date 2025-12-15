#!/usr/bin/env python3
"""
Unified Attachment Converter
모든 첨부파일을 텍스트로 변환하는 통합 변환기
지원: PDF, Word (docx/doc), HWP, 이미지 (OCR)
"""

import os
import sys
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, List, Union
import json
import tempfile
import shutil
from datetime import datetime
import argparse
import warnings
warnings.filterwarnings('ignore')

# PDF 변환 모듈 임포트
try:
    from .pdf_to_text import AdvancedPDFConverter
except ImportError:
    from pdf_to_text import AdvancedPDFConverter

# Word 문서 처리
try:
    from docx import Document  # python-docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client  # pywin32 (Windows에서 .doc 파일)
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

try:
    from doc2docx import convert as doc2docx_convert  # doc2docx
    DOC2DOCX_AVAILABLE = True
except ImportError:
    DOC2DOCX_AVAILABLE = False

# HWP 처리
try:
    import olefile
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False

try:
    from pyhwp import hwp5
    from pyhwp.hwp5.xmlmodel import Hwp5File
    PYHWP_AVAILABLE = True
except ImportError:
    PYHWP_AVAILABLE = False

# 이미지 OCR
try:
    from PIL import Image
    import pytesseract
    PIL_AVAILABLE = True
    TESSERACT_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    TESSERACT_AVAILABLE = False

# HTML 파싱 (HWP 변환용)
try:
    from bs4 import BeautifulSoup
    import html2text
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

# 추가 유틸리티
try:
    import magic  # python-magic (파일 타입 감지)
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False


class UnifiedAttachmentConverter:
    """통합 첨부파일 변환기"""

    def __init__(self, enable_ocr: bool = True, ocr_language: str = "kor+eng"):
        """
        초기화

        Args:
            enable_ocr: OCR 활성화 여부
            ocr_language: OCR 언어 설정
        """
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language

        # PDF 변환기 초기화 (선택적)
        try:
            self.pdf_converter = AdvancedPDFConverter(enable_ocr=enable_ocr)
        except ImportError as e:
            self.pdf_converter = None
            print(f"⚠️ PDF 라이브러리가 설치되지 않아 PDF 변환이 제한됩니다: {e}")

        self.supported_formats = self._get_supported_formats()

    def _get_supported_formats(self) -> Dict[str, List[str]]:
        """지원 가능한 포맷 확인"""
        formats = {
            "pdf": [] if self.pdf_converter is None else [".pdf"],
            "word": [],
            "hwp": [],
            "image": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif"],
            "text": [".txt", ".text", ".md", ".markdown", ".rst", ".log", ".csv", ".json", ".html", ".htm", ".xml"]
        }

        # Word 지원 확인
        if DOCX_AVAILABLE:
            formats["word"].append(".docx")
        if DOC_AVAILABLE or DOC2DOCX_AVAILABLE:
            formats["word"].append(".doc")

        # HWP 지원 확인
        if PYHWP_AVAILABLE or OLEFILE_AVAILABLE:
            formats["hwp"].extend([".hwp", ".hwpx", ".hml"])

        return formats

    def convert(self, file_path: str, output_format: str = "text") -> Dict[str, Any]:
        """
        파일을 텍스트로 변환 (자동 타입 감지)

        Args:
            file_path: 변환할 파일 경로
            output_format: 출력 형식 ("text", "json")

        Returns:
            변환 결과 딕셔너리
        """
        if not os.path.exists(file_path):
            return {
                "status": "error",
                "error": f"File not found: {file_path}",
                "file": file_path
            }

        # 파일 타입 감지
        file_type = self._detect_file_type(file_path)
        file_ext = Path(file_path).suffix.lower()

        print(f"\n📎 Attachment Converter")
        print(f"   File: {os.path.basename(file_path)}")
        print(f"   Type: {file_type} ({file_ext})")
        print(f"   Size: {os.path.getsize(file_path):,} bytes")

        result = None

        # 파일 타입별 처리
        if file_ext in self.supported_formats["pdf"]:
            result = self._convert_pdf(file_path)

        elif file_ext in self.supported_formats["word"]:
            result = self._convert_word(file_path)

        elif file_ext in self.supported_formats["hwp"]:
            result = self._convert_hwp(file_path)

        elif file_ext in self.supported_formats["image"]:
            result = self._convert_image_ocr(file_path)

        elif file_ext in self.supported_formats["text"]:
            result = self._convert_text(file_path)

        else:
            # 확장자로 판단 불가 시 MIME 타입으로 재시도
            if "pdf" in file_type.lower():
                result = self._convert_pdf(file_path)
            elif "word" in file_type.lower() or "document" in file_type.lower():
                result = self._convert_word(file_path)
            elif "image" in file_type.lower():
                result = self._convert_image_ocr(file_path)
            elif "text" in file_type.lower():
                result = self._convert_text(file_path)
            else:
                result = {
                    "status": "error",
                    "error": f"Unsupported file type: {file_type}",
                    "file": file_path,
                    "extension": file_ext
                }

        # 통계 추가
        if result and result.get("status") == "success":
            result["file_type"] = file_type  # file_type 추가
            result["summary"] = {
                "file": os.path.basename(file_path),
                "file_size": os.path.getsize(file_path),
                "conversion_method": result.get("method", "unknown"),
                "text_length": len(result.get("text", "")),
                "conversion_time": datetime.now().isoformat()
            }

        return result

    def _detect_file_type(self, file_path: str) -> str:
        """파일 타입 감지"""
        # python-magic 사용 (설치된 경우)
        if MAGIC_AVAILABLE:
            try:
                file_type = magic.from_file(file_path, mime=True)
                return file_type
            except:
                pass

        # 기본 mimetypes 사용
        mime_type, _ = mimetypes.guess_type(file_path)
        return mime_type or "unknown"

    def _convert_pdf(self, file_path: str) -> Dict[str, Any]:
        """PDF 변환"""
        print("   → PDF 변환 시작...")

        if self.pdf_converter is None:
            return {
                "status": "error",
                "error": "PDF converter not available. Install pypdf, pdfplumber, or pymupdf",
                "file": file_path
            }

        # 기존 PDF 변환기 사용
        result = self.pdf_converter.convert(file_path)

        if result.get("status") == "success":
            # 텍스트가 너무 적으면 OCR 시도
            if len(result.get("full_text", "").strip()) < 50 and self.enable_ocr:
                print("   → 텍스트 부족, OCR 시도...")
                ocr_result = self.pdf_converter.convert_with_ocr(file_path, self.ocr_language)
                if ocr_result.get("status") == "success":
                    result = ocr_result

            return {
                "status": "success",
                "method": result.get("method", "pdf"),
                "text": result.get("full_text", ""),
                "pages": result.get("total_pages", 0),
                "metadata": result.get("metadata", {}),
                "file": file_path
            }

        return result

    def _convert_word(self, file_path: str) -> Dict[str, Any]:
        """Word 문서 변환 (.docx, .doc)"""
        print("   → Word 문서 변환 시작...")

        file_ext = Path(file_path).suffix.lower()

        # .docx 파일 처리
        if file_ext == ".docx" and DOCX_AVAILABLE:
            return self._convert_docx(file_path)

        # .doc 파일 처리
        elif file_ext == ".doc":
            # 먼저 .docx로 변환 시도
            if DOC2DOCX_AVAILABLE:
                try:
                    temp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                    temp_docx.close()

                    print("   → .doc를 .docx로 변환 중...")
                    doc2docx_convert(file_path, temp_docx.name)

                    result = self._convert_docx(temp_docx.name)
                    os.unlink(temp_docx.name)
                    return result
                except Exception as e:
                    print(f"   ⚠️ doc2docx 변환 실패: {str(e)}")

            # Windows COM 사용
            if DOC_AVAILABLE and sys.platform == "win32":
                return self._convert_doc_win32(file_path)

            # 실패 시 기본 텍스트 추출
            return self._convert_as_text(file_path)

        return {
            "status": "error",
            "error": "Word conversion not available",
            "file": file_path
        }

    def _convert_docx(self, file_path: str) -> Dict[str, Any]:
        """DOCX 파일 변환"""
        try:
            doc = Document(file_path)
            full_text = []
            metadata = {}

            # 문서 속성 추출
            core_props = doc.core_properties
            metadata = {
                "title": getattr(core_props, 'title', ''),
                "author": getattr(core_props, 'author', ''),
                "subject": getattr(core_props, 'subject', ''),
                "keywords": getattr(core_props, 'keywords', ''),
                "created": str(getattr(core_props, 'created', '')),
                "modified": str(getattr(core_props, 'modified', ''))
            }

            # 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)

            # 테이블 텍스트 추출
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_data.append("\t".join(row_text))
                if table_data:
                    tables_text.append("\n".join(table_data))

            # 전체 텍스트 조합
            all_text = "\n\n".join(full_text)
            if tables_text:
                all_text += "\n\n[Tables]\n" + "\n\n".join(tables_text)

            return {
                "status": "success",
                "method": "docx",
                "text": all_text,
                "metadata": metadata,
                "paragraphs": len(full_text),
                "tables": len(doc.tables),
                "file": file_path
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file": file_path
            }

    def _convert_doc_win32(self, file_path: str) -> Dict[str, Any]:
        """Windows COM을 사용한 .doc 변환"""
        try:
            import pythoncom
            pythoncom.CoInitialize()

            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(os.path.abspath(file_path))
            text = doc.Range().Text

            doc.Close()
            word.Quit()

            return {
                "status": "success",
                "method": "doc_win32",
                "text": text,
                "file": file_path
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file": file_path
            }

    def _convert_hwp(self, file_path: str) -> Dict[str, Any]:
        """HWP 한글 파일 변환"""
        print("   → HWP 파일 변환 시작...")

        # pyhwp 사용 (권장)
        if PYHWP_AVAILABLE:
            try:
                return self._convert_hwp_pyhwp(file_path)
            except Exception as e:
                print(f"   ⚠️ pyhwp 실패: {str(e)}")

        # olefile로 기본 추출
        if OLEFILE_AVAILABLE:
            try:
                return self._convert_hwp_olefile(file_path)
            except Exception as e:
                print(f"   ⚠️ olefile 실패: {str(e)}")

        # 모두 실패 시
        return {
            "status": "error",
            "error": "HWP conversion not available. Install: pip install pyhwp olefile",
            "file": file_path
        }

    def _convert_hwp_pyhwp(self, file_path: str) -> Dict[str, Any]:
        """pyhwp를 사용한 HWP 변환"""
        try:
            from hwp5 import filestructure
            from hwp5.xmlmodel import Hwp5File

            hwp = Hwp5File(file_path)
            text_segments = []

            # BodyText 스트림에서 텍스트 추출
            for section in hwp.bodytext.sections:
                for paragraph in section:
                    para_text = []
                    for run in paragraph.runs:
                        if hasattr(run, 'text'):
                            para_text.append(run.text)
                    if para_text:
                        text_segments.append(''.join(para_text))

            full_text = '\n'.join(text_segments)

            # 문서 정보 추출
            metadata = {}
            if hasattr(hwp, 'summaryinfo'):
                summary = hwp.summaryinfo
                metadata = {
                    'title': getattr(summary, 'title', ''),
                    'author': getattr(summary, 'author', ''),
                    'subject': getattr(summary, 'subject', ''),
                    'keywords': getattr(summary, 'keywords', '')
                }

            return {
                "status": "success",
                "method": "pyhwp",
                "text": full_text,
                "metadata": metadata,
                "file": file_path
            }

        except Exception as e:
            raise e

    def _convert_hwp_olefile(self, file_path: str) -> Dict[str, Any]:
        """olefile을 사용한 기본 HWP 텍스트 추출"""
        try:
            import olefile
            import zlib

            ole = olefile.OleFileIO(file_path)
            text_segments = []

            # BodyText 섹션 탐색
            for stream_path in ole.listdir():
                if stream_path[0].startswith('BodyText'):
                    try:
                        data = ole.openstream(stream_path).read()
                        # 압축 해제 시도
                        try:
                            data = zlib.decompress(data, -15)
                        except:
                            pass

                        # UTF-16 디코딩 시도
                        try:
                            text = data.decode('utf-16-le', errors='ignore')
                            # 제어 문자 제거
                            text = ''.join(char for char in text if char.isprintable() or char.isspace())
                            if text.strip():
                                text_segments.append(text)
                        except:
                            # UTF-8 시도
                            try:
                                text = data.decode('utf-8', errors='ignore')
                                text = ''.join(char for char in text if char.isprintable() or char.isspace())
                                if text.strip():
                                    text_segments.append(text)
                            except:
                                pass
                    except:
                        continue

            ole.close()

            full_text = '\n'.join(text_segments)

            if not full_text.strip():
                return {
                    "status": "error",
                    "error": "No text extracted from HWP file",
                    "file": file_path
                }

            return {
                "status": "success",
                "method": "olefile",
                "text": full_text,
                "file": file_path
            }

        except Exception as e:
            raise e

    def _convert_image_ocr(self, file_path: str) -> Dict[str, Any]:
        """이미지 파일 OCR 변환"""
        print("   → 이미지 OCR 변환 시작...")

        if not (PIL_AVAILABLE and TESSERACT_AVAILABLE):
            return {
                "status": "error",
                "error": "OCR not available. Install: pip install pillow pytesseract",
                "file": file_path
            }

        if not self.enable_ocr:
            return {
                "status": "error",
                "error": "OCR is disabled",
                "file": file_path
            }

        try:
            # 이미지 열기
            image = Image.open(file_path)

            # 이미지 정보
            image_info = {
                "format": image.format,
                "mode": image.mode,
                "size": f"{image.width}x{image.height}",
                "dpi": image.info.get('dpi', 'unknown')
            }

            # OCR 수행
            print(f"   → OCR 언어: {self.ocr_language}")
            text = pytesseract.image_to_string(image, lang=self.ocr_language)

            # OCR 신뢰도
            ocr_data = pytesseract.image_to_data(image, lang=self.ocr_language, output_type=pytesseract.Output.DICT)
            confidence_scores = [c for c in ocr_data['conf'] if c > 0]
            avg_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0

            return {
                "status": "success",
                "method": "ocr",
                "text": text,
                "confidence": avg_confidence,
                "image_info": image_info,
                "language": self.ocr_language,
                "file": file_path
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file": file_path
            }

    def _convert_text(self, file_path: str) -> Dict[str, Any]:
        """텍스트 파일 읽기 (txt, csv, json, html 등)"""
        try:
            file_ext = Path(file_path).suffix.lower()

            # JSON 파일 처리
            if file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # JSON을 읽기 쉬운 텍스트로 변환
                    text = json.dumps(data, ensure_ascii=False, indent=2)
                    return {
                        "status": "success",
                        "method": "json",
                        "text": text,
                        "file": file_path
                    }

            # HTML 파일 처리
            elif file_ext in ['.html', '.htm']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                    # HTML 태그 제거
                    if BS4_AVAILABLE:
                        soup = BeautifulSoup(html_content, 'html.parser')
                        text = soup.get_text()
                    else:
                        # 간단한 태그 제거
                        import re
                        text = re.sub('<.*?>', '', html_content)
                    return {
                        "status": "success",
                        "method": "html",
                        "text": text.strip(),
                        "file": file_path
                    }

            # CSV 파일 처리
            elif file_ext == '.csv':
                import csv
                lines = []
                # 인코딩 자동 감지
                for encoding in ['utf-8', 'cp949', 'euc-kr', 'latin-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            reader = csv.reader(f)
                            for row in reader:
                                lines.append('\t'.join(row))
                        return {
                            "status": "success",
                            "method": "csv",
                            "text": '\n'.join(lines),
                            "encoding": encoding,
                            "file": file_path
                        }
                    except (UnicodeDecodeError, csv.Error):
                        continue

            # 일반 텍스트 파일
            else:
                # 인코딩 자동 감지
                encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']

                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                            return {
                                "status": "success",
                                "method": "text",
                                "text": text,
                                "encoding": encoding,
                                "file": file_path
                            }
                    except UnicodeDecodeError:
                        continue

            return {
                "status": "error",
                "error": "Failed to decode text file",
                "file": file_path
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file": file_path
            }

    def _convert_as_text(self, file_path: str) -> Dict[str, Any]:
        """최후의 수단: 바이너리를 텍스트로 변환 시도"""
        try:
            with open(file_path, 'rb') as f:
                data = f.read()

            # 다양한 인코딩 시도
            for encoding in ['utf-8', 'utf-16', 'cp949', 'euc-kr', 'latin-1']:
                try:
                    text = data.decode(encoding, errors='ignore')
                    # 읽을 수 있는 텍스트만 필터링
                    text = ''.join(char for char in text if char.isprintable() or char.isspace())

                    if len(text.strip()) > 50:  # 최소 50자 이상
                        return {
                            "status": "success",
                            "method": "fallback_text",
                            "text": text,
                            "encoding": encoding,
                            "file": file_path
                        }
                except:
                    continue

            return {
                "status": "error",
                "error": "Unable to extract text",
                "file": file_path
            }

        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "file": file_path
            }

    def batch_convert(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """여러 파일 일괄 변환"""
        results = []
        total = len(file_paths)

        print(f"\n📎 Batch Conversion: {total} files")

        for idx, file_path in enumerate(file_paths, 1):
            print(f"\n[{idx}/{total}] Processing...")
            result = self.convert(file_path)
            results.append(result)

            if result.get("status") == "success":
                text_len = len(result.get("text", ""))
                print(f"   ✅ Success: {text_len:,} characters extracted")
            else:
                print(f"   ❌ Failed: {result.get('error', 'Unknown error')}")

        # 요약 통계
        success_count = sum(1 for r in results if r.get("status") == "success")
        print(f"\n📊 Results: {success_count}/{total} successful")

        return results

    def save_result(self, result: Dict[str, Any], output_path: Optional[str] = None, format: str = "text") -> str:
        """변환 결과 저장"""
        if result.get("status") != "success":
            raise ValueError(f"Conversion failed: {result.get('error')}")

        if output_path is None:
            base_name = Path(result["file"]).stem
            output_path = f"{base_name}_converted.{format}"

        if format == "json":
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
        else:  # text
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"Attachment Conversion Result\n")
                f.write(f"{'=' * 80}\n")
                f.write(f"Source: {result['file']}\n")
                f.write(f"Method: {result.get('method', 'unknown')}\n")
                f.write(f"Time: {datetime.now().isoformat()}\n")

                if result.get('metadata'):
                    f.write(f"\nMetadata:\n")
                    for key, value in result['metadata'].items():
                        if value:
                            f.write(f"  {key}: {value}\n")

                f.write(f"\n{'=' * 80}\n\n")
                f.write(result.get('text', ''))

        print(f"💾 Saved to: {output_path}")
        return output_path


class AttachmentAPI:
    """외부 사용을 위한 간단한 API"""

    def __init__(self):
        self.converter = UnifiedAttachmentConverter(enable_ocr=True)

    def convert_to_text(self, file_path: str) -> str:
        """
        파일을 텍스트로 변환 (단순 인터페이스)

        Args:
            file_path: 파일 경로

        Returns:
            추출된 텍스트 문자열

        Raises:
            Exception: 변환 실패 시
        """
        result = self.converter.convert(file_path)

        if result.get("status") == "success":
            return result.get("text", "")
        else:
            raise Exception(f"Conversion failed: {result.get('error', 'Unknown error')}")

    def convert_with_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        파일을 변환하고 메타데이터 포함

        Args:
            file_path: 파일 경로

        Returns:
            변환 결과와 메타데이터

        Raises:
            Exception: 변환 실패 시
        """
        result = self.converter.convert(file_path)

        if result.get("status") == "success":
            return {
                "text": result.get("text", ""),
                "metadata": result.get("metadata", {}),
                "method": result.get("method", ""),
                "summary": result.get("summary", {})
            }
        else:
            raise Exception(f"Conversion failed: {result.get('error', 'Unknown error')}")

    def check_support(self, file_path: str) -> bool:
        """
        파일 지원 여부 확인

        Args:
            file_path: 파일 경로

        Returns:
            지원 여부
        """
        ext = Path(file_path).suffix.lower()

        for format_list in self.converter.supported_formats.values():
            if ext in format_list:
                return True

        return False

    def get_supported_formats(self) -> Dict[str, List[str]]:
        """지원 포맷 목록 반환"""
        return self.converter.supported_formats


def main():
    """CLI 인터페이스"""
    parser = argparse.ArgumentParser(
        description="Unified Attachment Converter",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  %(prog)s document.pdf              # PDF 변환
  %(prog)s report.docx               # Word 변환
  %(prog)s document.hwp              # HWP 변환
  %(prog)s scan.jpg --ocr            # 이미지 OCR
  %(prog)s *.* -d output/            # 모든 파일 일괄 변환

Supported Formats:
  - PDF: .pdf (자동 OCR 지원)
  - Word: .docx, .doc
  - HWP: .hwp, .hwpx
  - Images: .jpg, .png, .tiff, etc. (OCR)
  - Text: .txt, .md, .log
        """
    )

    parser.add_argument("files", nargs="+", help="File(s) to convert")
    parser.add_argument("-o", "--output", help="Output file path")
    parser.add_argument("-f", "--format", choices=["text", "json"], default="text",
                       help="Output format (default: text)")
    parser.add_argument("-d", "--directory", help="Output directory for batch")
    parser.add_argument("--no-ocr", action="store_true", help="Disable OCR")
    parser.add_argument("-l", "--language", default="kor+eng",
                       help="OCR language (default: kor+eng)")
    parser.add_argument("--check", action="store_true",
                       help="Check if files are supported")
    parser.add_argument("--api-demo", action="store_true",
                       help="Show API usage example")

    args = parser.parse_args()

    # API 데모
    if args.api_demo:
        print("\n" + "=" * 80)
        print("API Usage Example:")
        print("=" * 80)
        print("""
from attachment_converter import AttachmentAPI

# API 초기화
api = AttachmentAPI()

# 간단한 텍스트 추출
text = api.convert_to_text("document.pdf")
print(text)

# 메타데이터 포함 추출
result = api.convert_with_metadata("report.docx")
print(f"Text: {result['text'][:100]}...")
print(f"Method: {result['method']}")
print(f"Metadata: {result['metadata']}")

# 지원 여부 확인
if api.check_support("file.hwp"):
    print("Supported!")

# 지원 포맷 목록
formats = api.get_supported_formats()
print(formats)
        """)
        return

    # 변환기 초기화
    converter = UnifiedAttachmentConverter(
        enable_ocr=not args.no_ocr,
        ocr_language=args.language
    )

    # 지원 여부만 확인
    if args.check:
        print("\n📋 Checking file support...")
        for file_path in args.files:
            ext = Path(file_path).suffix.lower()
            supported = False

            for format_type, extensions in converter.supported_formats.items():
                if ext in extensions:
                    supported = True
                    print(f"  ✅ {file_path} - Supported ({format_type})")
                    break

            if not supported:
                print(f"  ❌ {file_path} - Not supported")
        return

    # 단일 파일 변환
    if len(args.files) == 1:
        file_path = args.files[0]
        result = converter.convert(file_path)

        if result.get("status") == "success":
            output_path = converter.save_result(result, args.output, args.format)
            print(f"\n✅ Conversion successful!")
            print(f"   Characters: {len(result.get('text', '')):,}")
            print(f"   Output: {output_path}")
        else:
            print(f"\n❌ Conversion failed: {result.get('error')}")
            return 1

    # 일괄 변환
    else:
        output_dir = args.directory or "."
        os.makedirs(output_dir, exist_ok=True)

        results = converter.batch_convert(args.files)

        for result in results:
            if result.get("status") == "success":
                try:
                    base_name = Path(result["file"]).stem
                    output_path = os.path.join(output_dir, f"{base_name}.{args.format}")
                    converter.save_result(result, output_path, args.format)
                except Exception as e:
                    print(f"   Failed to save: {e}")


if __name__ == "__main__":
    main()