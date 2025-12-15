#!/usr/bin/env python3
"""
Simple Attachment Converter - 간단한 인터페이스
사용자가 쉽게 사용할 수 있는 간단한 API
"""

import os
import json
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any, Union, List
import tempfile
import urllib.request
import urllib.parse
from datetime import datetime


def convert_to_text(file_path_or_url: str, **kwargs) -> str:
    """
    파일이나 URL을 텍스트로 변환

    Args:
        file_path_or_url: 로컬 파일 경로 또는 URL
        **kwargs: 추가 옵션 (encoding, ocr_language 등)

    Returns:
        추출된 텍스트 문자열

    Examples:
        >>> text = convert_to_text("document.pdf")
        >>> text = convert_to_text("https://example.com/file.pdf")
        >>> text = convert_to_text("korean.txt", encoding="cp949")
    """
    converter = SimpleConverter(**kwargs)
    return converter.convert(file_path_or_url)


def extract_text(file_path: str, **kwargs) -> str:
    """convert_to_text의 별칭 (더 직관적인 이름)"""
    return convert_to_text(file_path, **kwargs)


def batch_convert(file_list: List[str], **kwargs) -> Dict[str, str]:
    """
    여러 파일을 한 번에 변환

    Args:
        file_list: 파일 경로 리스트
        **kwargs: 추가 옵션

    Returns:
        {파일경로: 텍스트} 딕셔너리

    Example:
        >>> texts = batch_convert(["file1.pdf", "file2.txt", "file3.html"])
    """
    converter = SimpleConverter(**kwargs)
    results = {}

    for file_path in file_list:
        try:
            results[file_path] = converter.convert(file_path)
        except Exception as e:
            results[file_path] = f"Error: {str(e)}"

    return results


def is_supported(file_path: str) -> bool:
    """
    파일이 지원되는지 확인

    Args:
        file_path: 파일 경로

    Returns:
        지원 여부

    Example:
        >>> if is_supported("document.pdf"):
        ...     text = convert_to_text("document.pdf")
    """
    ext = Path(file_path).suffix.lower()
    supported = {
        '.pdf', '.txt', '.text', '.md', '.markdown', '.rst', '.log',
        '.html', '.htm', '.xml', '.json', '.csv', '.tsv',
        '.docx', '.doc', '.hwp', '.hwpx',
        '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif'
    }
    return ext in supported


class SimpleConverter:
    """내부 구현 클래스 (사용자는 직접 사용할 필요 없음)"""

    def __init__(self, encoding: str = 'auto', ocr_language: str = 'eng+kor', **kwargs):
        self.encoding = encoding
        self.ocr_language = ocr_language
        self.options = kwargs

    def convert(self, source: str) -> str:
        """파일 또는 URL을 텍스트로 변환"""

        # URL인지 확인
        if source.startswith(('http://', 'https://', 'ftp://')):
            return self._convert_url(source)

        # 로컬 파일
        if not os.path.exists(source):
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {source}")

        # 파일 확장자로 타입 결정
        ext = Path(source).suffix.lower()

        # PDF
        if ext == '.pdf':
            return self._convert_pdf(source)

        # Word
        elif ext in ['.docx', '.doc']:
            return self._convert_word(source)

        # HWP
        elif ext in ['.hwp', '.hwpx']:
            return self._convert_hwp(source)

        # 이미지
        elif ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif']:
            return self._convert_image(source)

        # HTML/XML
        elif ext in ['.html', '.htm', '.xml']:
            return self._convert_html(source)

        # JSON
        elif ext == '.json':
            return self._convert_json(source)

        # CSV/TSV
        elif ext in ['.csv', '.tsv']:
            return self._convert_csv(source)

        # 일반 텍스트
        else:
            return self._convert_text(source)

    def _convert_url(self, url: str) -> str:
        """URL에서 파일 다운로드 후 변환"""
        try:
            # URL에서 파일 확장자 추측
            parsed = urllib.parse.urlparse(url)
            path = parsed.path
            ext = Path(path).suffix.lower() or '.tmp'

            # 임시 파일로 다운로드
            with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp_file:
                urllib.request.urlretrieve(url, tmp_file.name)
                temp_path = tmp_file.name

            try:
                # 다운로드된 파일 변환
                return self.convert(temp_path)
            finally:
                # 임시 파일 삭제
                os.unlink(temp_path)

        except Exception as e:
            raise Exception(f"URL 처리 실패: {str(e)}")

    def _convert_pdf(self, file_path: str) -> str:
        """PDF 변환"""
        try:
            # pypdf 시도
            try:
                from pypdf import PdfReader

                reader = PdfReader(file_path)
                text_parts = []

                for page_num, page in enumerate(reader.pages, 1):
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{text}")

                return '\n\n'.join(text_parts)

            except ImportError:
                pass

            # PyMuPDF 시도
            try:
                import fitz

                doc = fitz.open(file_path)
                text_parts = []

                for page_num, page in enumerate(doc, 1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num} ---\n{text}")

                doc.close()
                return '\n\n'.join(text_parts)

            except ImportError:
                pass

            # pdfplumber 시도
            try:
                import pdfplumber

                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"--- Page {page_num} ---\n{text}")

                return '\n\n'.join(text_parts)

            except ImportError:
                pass

            raise ImportError("PDF 라이브러리가 설치되지 않았습니다. pip install pypdf")

        except Exception as e:
            raise Exception(f"PDF 변환 실패: {str(e)}")

    def _convert_word(self, file_path: str) -> str:
        """Word 문서 변환"""
        try:
            from docx import Document

            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            # 테이블 텍스트도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            return '\n\n'.join(paragraphs)

        except ImportError:
            # .doc 파일이거나 python-docx 없는 경우 기본 텍스트 추출
            return self._extract_basic_text(file_path)
        except Exception as e:
            raise Exception(f"Word 변환 실패: {str(e)}")

    def _convert_hwp(self, file_path: str) -> str:
        """HWP 한글 문서 변환"""
        try:
            import olefile
            import zlib

            ole = olefile.OleFileIO(file_path)
            text_parts = []

            for entry in ole.listdir():
                if entry[0].startswith('BodyText'):
                    data = ole.openstream(entry).read()

                    # 압축 해제 시도
                    try:
                        data = zlib.decompress(data, -15)
                    except:
                        pass

                    # 텍스트 디코딩
                    for encoding in ['utf-16-le', 'utf-8', 'cp949']:
                        try:
                            text = data.decode(encoding, errors='ignore')
                            text = ''.join(c for c in text if c.isprintable() or c.isspace())
                            if text.strip():
                                text_parts.append(text)
                                break
                        except:
                            continue

            ole.close()
            return '\n'.join(text_parts)

        except ImportError:
            raise ImportError("HWP 변환을 위해 olefile 설치 필요: pip install olefile")
        except Exception as e:
            raise Exception(f"HWP 변환 실패: {str(e)}")

    def _convert_image(self, file_path: str) -> str:
        """이미지 OCR 변환"""
        try:
            from PIL import Image
            import pytesseract

            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang=self.ocr_language)
            return text

        except ImportError:
            raise ImportError("OCR을 위해 설치 필요: pip install pillow pytesseract")
        except Exception as e:
            raise Exception(f"이미지 OCR 실패: {str(e)}")

    def _convert_html(self, file_path: str) -> str:
        """HTML/XML 변환"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # BeautifulSoup 있으면 사용
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')
            return soup.get_text()
        except ImportError:
            # 없으면 간단한 태그 제거
            import re
            text = re.sub('<[^<]+?>', '', content)
            return text.strip()

    def _convert_json(self, file_path: str) -> str:
        """JSON 파일 변환"""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # JSON을 읽기 쉬운 텍스트로
        return json.dumps(data, ensure_ascii=False, indent=2)

    def _convert_csv(self, file_path: str) -> str:
        """CSV/TSV 파일 변환"""
        import csv

        # 구분자 자동 감지
        with open(file_path, 'r', encoding=self._detect_encoding(file_path)) as f:
            sample = f.read(1024)
            delimiter = '\t' if '\t' in sample else ','
            f.seek(0)

            reader = csv.reader(f, delimiter=delimiter)
            rows = ['\t'.join(row) for row in reader]

        return '\n'.join(rows)

    def _convert_text(self, file_path: str) -> str:
        """일반 텍스트 파일 변환"""
        encoding = self.encoding

        if encoding == 'auto':
            encoding = self._detect_encoding(file_path)

        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()

    def _detect_encoding(self, file_path: str) -> str:
        """텍스트 파일 인코딩 자동 감지"""
        encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1', 'utf-16']

        for enc in encodings:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    f.read()
                return enc
            except UnicodeDecodeError:
                continue

        return 'utf-8'  # 기본값

    def _extract_basic_text(self, file_path: str) -> str:
        """바이너리 파일에서 기본 텍스트 추출"""
        with open(file_path, 'rb') as f:
            data = f.read()

        # 다양한 인코딩으로 시도
        for encoding in ['utf-8', 'utf-16', 'cp949', 'latin-1']:
            try:
                text = data.decode(encoding, errors='ignore')
                # 읽을 수 있는 문자만 필터
                text = ''.join(c for c in text if c.isprintable() or c.isspace())
                if len(text.strip()) > 50:
                    return text
            except:
                continue

        return ""


# 간단한 사용 예제를 위한 헬퍼 함수들

def extract_pdf_text(pdf_path: str) -> str:
    """PDF에서 텍스트 추출 (특화 함수)"""
    return convert_to_text(pdf_path)


def extract_word_text(word_path: str) -> str:
    """Word에서 텍스트 추출 (특화 함수)"""
    return convert_to_text(word_path)


def extract_from_url(url: str) -> str:
    """URL에서 파일 다운로드 후 텍스트 추출"""
    return convert_to_text(url)


def quick_convert(*files) -> Union[str, Dict[str, str]]:
    """
    빠른 변환 - 하나 또는 여러 파일 처리

    Example:
        >>> text = quick_convert("document.pdf")
        >>> texts = quick_convert("doc1.pdf", "doc2.txt", "doc3.html")
    """
    if len(files) == 1:
        return convert_to_text(files[0])
    else:
        return batch_convert(list(files))


# __all__을 정의하여 공개 API 명시
__all__ = [
    'convert_to_text',
    'extract_text',
    'batch_convert',
    'is_supported',
    'extract_pdf_text',
    'extract_word_text',
    'extract_from_url',
    'quick_convert'
]


if __name__ == "__main__":
    # 간단한 테스트
    import sys

    if len(sys.argv) < 2:
        print("사용법: python simple_converter.py <파일경로>")
        print("\n예제:")
        print("  python simple_converter.py document.pdf")
        print("  python simple_converter.py https://example.com/file.pdf")
        sys.exit(1)

    try:
        text = convert_to_text(sys.argv[1])
        print(f"추출된 텍스트 ({len(text)} 문자):")
        print("-" * 50)
        print(text[:500] + "..." if len(text) > 500 else text)
    except Exception as e:
        print(f"오류: {e}")